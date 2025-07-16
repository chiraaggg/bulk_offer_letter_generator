import streamlit as st
import pandas as pd
from docx import Document
from docx2pdf import convert
import os
from zipfile import ZipFile
from datetime import datetime
import shutil

st.set_page_config(page_title="Word Offer Letter Generator")
st.title("📄 Word → PDF Offer Letter Generator")

template_docx = st.file_uploader("Upload Word Template (.docx)", type=["docx"])
data_file = st.file_uploader("Upload Candidate Excel or CSV", type=["xlsx", "csv"])

output_dir = "output_letters"
os.makedirs(output_dir, exist_ok=True)
temp_word_dir = "temp_word_files"
os.makedirs(temp_word_dir, exist_ok=True)

from docx import Document

def fmt(val):
    return f"₹{val:,.2f}"

from docx import Document

def fmt(val):
    return f"₹{val:,.2f}"

def generate_docx(template_path, data_row, output_path):
    # Step 1: Calculate CTC
    annual_ctc = float(data_row["Annual CTC"])
    monthly_ctc = annual_ctc / 12

    basicm = monthly_ctc * 0.5
    hram = basicm * 0.5
    splm = monthly_ctc - basicm - hram
    grossm = monthly_ctc
    ctcm = monthly_ctc

    basica = annual_ctc * 0.5
    hraa = basica * 0.5
    spla = annual_ctc - basica - hraa
    grossa = annual_ctc
    ctca = annual_ctc

    # Step 2: Merge values
    calc_fields = {
        "basicm": fmt(basicm),
        "hram": fmt(hram),
        "splm": fmt(splm),
        "grossm": fmt(grossm),
        "ctcm": fmt(ctcm),
        "basica": fmt(basica),
        "hraa": fmt(hraa),
        "spla": fmt(spla),
        "grossa": fmt(grossa),
        "ctca": fmt(ctca),
    }

    merged_fields = {**{k: str(v) for k, v in data_row.items()}, **calc_fields}

    # Step 3: Load document
    doc = Document(template_path)

    # Replace in paragraphs
    for p in doc.paragraphs:
        for key, val in merged_fields.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, val)

    # ✅ Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in merged_fields.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, val)

    doc.save(output_path)



if template_docx and data_file:
    ext = data_file.name.split(".")[-1]
    df = pd.read_excel(data_file) if ext == "xlsx" else pd.read_csv(data_file)

    if st.button("Generate Offer Letters"):
        st.info("Processing...")

        template_path = os.path.join(temp_word_dir, "template.docx")
        with open(template_path, "wb") as f:
            f.write(template_docx.read())

        pdf_paths = []
        for i, row in df.iterrows():
            name_slug = row["candidate full name"].replace(" ", "_")
            filled_docx = os.path.join(temp_word_dir, f"{name_slug}.docx")
            filled_pdf = os.path.join(output_dir, f"{name_slug}.pdf")

            # Replace placeholders
            generate_docx(template_path, row.to_dict(), filled_docx)

            # Convert to PDF
            convert(filled_docx, filled_pdf)
            pdf_paths.append(filled_pdf)

        # Create zip
        zip_path = os.path.join(output_dir, "Offer_Letters.zip")
        with ZipFile(zip_path, "w") as zipf:
            for pdf in pdf_paths:
                zipf.write(pdf, arcname=os.path.basename(pdf))

        with open(zip_path, "rb") as f:
            st.success("✅ Done! Download your ZIP:")
            st.download_button("📥 Download ZIP", f, "Offer_Letters.zip", "application/zip")

        # Clean temp files
        shutil.rmtree(temp_word_dir)
